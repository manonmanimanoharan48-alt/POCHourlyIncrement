#!/usr/bin/env python
# coding: utf-8

# ## POCHourlyIncrementNotebook
# 
# null

# ##### **Install country and docs library.**

# In[1]:


# The command is not a standard IPython magic command. It is designed for use within Fabric notebooks only.
# %pip install pycountry
# The command is not a standard IPython magic command. It is designed for use within Fabric notebooks only.
# %pip install python-docx


# In[2]:


# Install / upgrade GX properly
# The command is not a standard IPython magic command. It is designed for use within Fabric notebooks only.
# %pip install great-expectations==0.18.12 typing_extensions==4.12.2
# The command is not a standard IPython magic command. It is designed for use within Fabric notebooks only.
# %pip install pandas pyarrow


# 

# ##### **Fabric Setup Cell for Spark Session and Great Expectations**

# In[3]:


from pyspark.sql import SparkSession
from pyspark.sql.functions import *
from pyspark.sql.types import *
from pyspark.sql import functions as F
from notebookutils import variableLibrary
from delta.tables import DeltaTable
from datetime import datetime
import requests, json
import pycountry
import os

# Initialize spark session.
try:
    spark = SparkSession.builder \
        .appName("WeatherIncrementalHourlyAgg") \
        .getOrCreate()
except Exception as e:
    raise RuntimeError("Unable to get Spark session for weather incremental aggregation: " + str(e))
#Lakehouse paths for bronze and watermark table.
bronze_path = "Tables/bronze/weather_data"
watermark_table_path = "Tables/bronze/watermark_table"

# Get the variable library
config_lib = variableLibrary.getLibrary("POCVariableLibrary")


# ##### **FabricLogger Utility Class**

# In[4]:


class FabricLogger:
    """
    A production-ready logger for Microsoft Fabric PySpark notebooks.
    ✅ Writes INFO and ERROR logs
    ✅ Rotates logs daily
    ✅ Allows reading and filtering logs
    """

    def __init__(self, base_log_dir="/lakehouse/default/Files/logs/"):
        self.base_log_dir = base_log_dir
        os.makedirs(self.base_log_dir, exist_ok=True)
        self.log_file = self._get_daily_log_file()

    def _get_daily_log_file(self):
        """Generate today's log file name dynamically."""
        today = datetime.now().strftime("%Y-%m-%d")
        return os.path.join(self.base_log_dir, f"fabric_log_{today}.txt")

    def _write_log(self, level, message):
        """Internal: Write a log entry to today’s log file."""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        entry = f"[{timestamp}] [{level}] {message}\n"

        # Append entry to today’s log
        with open(self._get_daily_log_file(), "a") as f:
            f.write(entry)

    def info(self, message):
        """Write an INFO log."""
        self._write_log("INFO", message)

    def error(self, message):
        """Write an ERROR log."""
        self._write_log("ERROR", message)

    def read_logs(self, date=None, level_filter=None, tail=None):
        """
        Read log entries.
        - date: specific date in YYYY-MM-DD format (default = today)
        - level_filter: filter by 'INFO' or 'ERROR'
        - tail: number of last lines to show
        """
        if not date:
            date = datetime.now().strftime("%Y-%m-%d")

        log_file = os.path.join(self.base_log_dir, f"fabric_log_{date}.txt")

        if not os.path.exists(log_file):
            print(f"⚠️ No log file found for {date}")
            return

        with open(log_file, "r") as f:
            lines = f.readlines()

        if level_filter:
            lines = [l for l in lines if f"[{level_filter}]" in l]

        if tail:
            lines = lines[-tail:]

        print(f"📄 Logs from {log_file}:")
        for line in lines:
            print(line.strip())

    def list_log_files(self):
        """List all log files in the directory."""
        files = sorted(os.listdir(self.base_log_dir))
        for f in files:
            print(f)


# ##### **Fetch New Weather Data From Weather Gloabal API**

# In[5]:


# Initialize logger
logger = FabricLogger()

# Access API key from variable library.
API_KEY = config_lib.hourlyIncrementDevAPIKey
print(f"API_KEY -->{API_KEY}")

#API_KEY = "f20f64c78f0f45ba47620065f236a6d5"
cities = ["London", "New York", "Tokyo", "Sydney", "Mumbai"]

def fetch_weather(city, max_retries=3):
        # Construct the API URL with city and API key formatted for geo json output.
        url = f"http://api.openweathermap.org/data/2.5/weather?q={city}&appid={API_KEY}&units=metric"
        for attempt in range(1, max_retries):
            try:
                # Make the GET request to fetch data
                response = requests.get(url)
                # Check if the request was successful
                if response.status_code == 200:
                    data = response.json()
                    logger.info(f"Weather data successfully fetched for city :{city} and attemps is {attempt}")
                    return {
                        "city": city,
                        "temperature": data["main"]["temp"],
                        "humidity": data["main"]["humidity"],
                        "timestamp": datetime.utcfromtimestamp(data["dt"]).strftime("%Y-%m-%d %H:%M:%S"),
                        "load_time": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),  # for incremental control
                        "description": data['weather'][0]['description'],
                        "sunrise": datetime.utcfromtimestamp(data["sys"]["sunrise"]).strftime("%Y-%m-%d %H:%M:%S"),
                        "windspeed": data["wind"]["speed"],
                        "timezone" : data["timezone"],
                        "countrycode" :data["sys"]["country"]
                    }
                else:
                    logger.error(f"Failed to fetch weather data. Status code:{city} , response code: {response.status_code}")
                    #To stop the SparkSession.
                    mssparkutils.session.stop()
                    return {}
            except Exception as e:
                logger.info(f"Attempt {attempt + 1} failed: {e}")
                if attempt < max_retries - 1:
                    delay = initial_delay * (backoff_factor ** attempt)
                    logger.info(f"Retrying in {delay} seconds...")
                    time.sleep(delay)
                else:
                    logger.error("Max retries reached. Giving up.")
                    raise  # Re-raise the last exception if all retries fail
new_data = [fetch_weather(c) for c in cities]

#create schema for dataframe
dfSchema = StructType([StructField("city", StringType(), True)\
                    ,StructField("temperature", DoubleType(), True)\
                        ,StructField("timestamp", StringType(), True)\
                        ,StructField("load_time", StringType(), True)\
                        ,StructField("description", StringType(), True)\
                        ,StructField("sunrise", StringType(), True)\
                        ,StructField("windspeed", DoubleType(), True)\
                        ,StructField("timezone", LongType(), True)\
                        ,StructField("countrycode", StringType(), True)\
                        ,StructField("humidity", IntegerType(), True)])
new_df = spark.createDataFrame(new_data, schema = dfSchema)
   
# Convert offset in seconds → UTC±HH:MM
new_df = new_df.withColumn(
    "timezone_readable",
    F.format_string(
        "UTC%+03d:%02d",
        (F.col("timezone") / 3600).cast("int"),
        (F.abs(F.col("timezone") % 3600) / 60).cast("int")
    )
)
new_df = new_df.withColumn("event_time", to_timestamp(col("timestamp")))
new_df= new_df.drop("timeZone")
new_df= new_df.withColumnRenamed("timezone_readable", "timeZone")

# Create Python dictionary mapping
country_dict = {c.alpha_2: c.name for c in pycountry.countries}

# Broadcast mapping to all Spark workers
broadcast_country_map = spark.sparkContext.broadcast(country_dict)
# Define a UDF to map code → name
@F.udf("string")
def get_country_name(code):
    return broadcast_country_map.value.get(code, "Unknown")

# Apply it
new_df = new_df.withColumn("country", get_country_name(F.col("countrycode")))


# ##### **Incremental Data Load into Bronze (Delta Table)**

# In[6]:


# Check if the Bronze table exists
if DeltaTable.isDeltaTable(spark, bronze_path):
    bronze_tbl = DeltaTable.forPath(spark, bronze_path)
    wat_mark_tb= DeltaTable.forPath(spark, watermark_table_path)
        
    # Find max timestamp in existing watermark table.
    last_watermark = wat_mark_tb.toDF().filter(col("TableName") == 'weather_data').select("WatermarkValue").collect()[0][0]
    print(f"Last watermark -->: {last_watermark}")
    incremental_df = new_df.filter(col("timestamp") > last_watermark)
else:
    incremental_df = new_df

# Add partition columns -> only add partition to more than 10 records.
incremental_df = (
    incremental_df
        .withColumn("year", year("timestamp"))
        .withColumn("month", month("timestamp"))
)

try:
    # Append new data
    incremental_df.write.format("delta").partitionBy("year", "month").mode("append").save(bronze_path)
    logger.info(f"Data appended successfully in bronze path.")
except Exception as e:
    logger.error(f"Error occurred while appending data into bronze path: {str(e)}")

# Update watermark
new_watermark = incremental_df.agg({"timestamp": "max"}).collect()[0][0]
if new_watermark is None:
    logger.info(f"Latest watermark is empty.")
    print("Latest watermark is empty.")
else:
    logger.info(f"Latest watermark -->{new_watermark}")
    spark.sql(f"UPDATE bronze.watermark_table SET WatermarkValue = '{new_watermark}' WHERE TableName = 'weather_data'")


# ##### **Bronze Layer validation With GX**

# In[7]:


from pyspark.sql.functions import round, col
from great_expectations.data_context import get_context
from great_expectations.dataset import SparkDFDataset
import pandas as pd
import json
from docx import Document
from docx.shared import Pt

# Create a GX Ephemeral context (in-memory)
context = get_context(mode="ephemeral")
bronze_path = "Tables/bronze/weather_data"
# Expectation parameters (example)
BRONZE_EXPECTED_COLUMNS = ["city","temperature", "timestamp", "load_time","description","sunrise",
                    "windspeed","countrycode","humidity", "timeZone", "event_time", "country", "year", "month"]
BRONZE_NOT_NULL_COLUMNS =  ["city","temperature", "timestamp", "load_time","description","sunrise",
                    "windspeed","countrycode","humidity", "timeZone", "event_time", "country"]
BRONZE_EXIST_COLUMNS =  ["city","temperature", "timestamp", "load_time","description","sunrise",
                    "windspeed","countrycode","humidity", "timeZone", "event_time", "country"]
BRONZE_DATATYPE_CHECKS = {"city":"StringType","temperature": "DoubleType", "timestamp": "StringType", "load_time": "StringType",
 "description": "StringType", "sunrise": "StringType", "windspeed": "DoubleType", "countrycode": "StringType",
  "humidity": "IntegerType", "timeZone": "StringType", "event_time": "TimestampType", "country": "StringType"}
BRONZE_DATETIME_CHECKS = {"timestamp":"%Y-%m-%d %H:%M:%S", "load_time":"%Y-%m-%d %H:%M:%S", "sunrise":"%Y-%m-%d %H:%M:%S"}
BRONZE_RANGE_CHECKS = {"temperature": {"min_value": 0, "max_value": 50}}
BRONZE_REGEX_CHECKS = {"load_time": r"(^\d{4}-(0[1-9]|1[0-2])-(0[1-9]|[12]\d|3[01]) ([01]\d|2[0-3]):[0-5]\d:[0-5]\d$)|(^\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}Z$)",
                       "timeZone":r"(^UTC[+-](0[0-9]|1[0-4]):[0-5][0-9]$)",
                       "countrycode":r"(^[a-zA-Z]{2}$)"}


# Utility function for show passed/failed expectations.
def show_failed_expectations(results):
    failed = [r for r in results["results"] if not r["success"]]
    if not failed:
        print("✅ All expectations passed!")
    else:
        print("❌ Failed expectations:")
        for f in failed:
            print(f"- {f['expectation_config']['expectation_type']} failed")
def extract_stats(name, gx_result):
    data = gx_result.to_json_dict()["statistics"]
    return {
        "layer": name,
        "success_percent": data["success_percent"],
        "total": data["evaluated_expectations"],
        "passed": data["successful_expectations"],
        "failed": data["unsuccessful_expectations"]
    }
def write_test_summary_on_doc(layer_name, stats_bronze, summary_json):
     # Create validation folder in Files/
     validation_report_path = "/lakehouse/default/Files/validation_reports"
     os.makedirs(validation_report_path, exist_ok=True)
     report_file = f"{validation_report_path}/Data_Validation_Report.docx"
     # If the report exists, open it; otherwise, create a new one
     if os.path.exists(report_file):
        doc = Document(report_file)
     else:
        doc = Document()
        doc.add_heading("Data Validation Report (Microsoft Fabric + Great Expectations)", level=1)
     
     # Section: Bronze Layer or Silver Layer.
     doc.add_heading(f"{layer_name} Layer Validation", level=2)

     # Timestamped run header
     run_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
     doc.add_paragraph(f"\n Validation Run: {run_time}", style="Intense Quote")
     #stats_bronze = bronze_json["statistics"]

     doc.add_paragraph(f"✅ Success: {stats_bronze['success_percent']}%")
     doc.add_paragraph(f"Total Expectations: {stats_bronze['total']}")
     doc.add_paragraph(f"Passed: {stats_bronze['passed']}")
     doc.add_paragraph(f"Failed: {stats_bronze['failed']}")

     # Add detailed failed expectations
     doc.add_heading(f"{layer_name} Failed Expectations", level=3)
     failed_layer = [r for r in summary_json["results"] if not r["success"]]
     if not failed_layer:
        doc.add_paragraph("All expectations passed successfully.")
     else:
        for f in failed_layer:
            exp_type = f["expectation_config"]["expectation_type"]
            doc.add_paragraph(f"- {exp_type}", style="List Bullet")
    
     # Save report to Fabric Lakehouse Files folder
     os.makedirs(validation_report_path, exist_ok=True)
     doc.save(report_file)
     print(f"📄 {layer_name} Validation report saved successfully at: {report_file}")
     # Save both reports(json and docs file foramt)
     with open(f"{validation_report_path}/{layer_name}_validation.json", "w") as f:
        json.dump(summary_json, f, indent=2)
     print(f"📁 Reports saved under: {validation_report_path}")
try:
    bronze_df = spark.read.format("delta").load(bronze_path)
    # Convert to GX Dataset
    gx_bronze = SparkDFDataset(bronze_df)

    # Define expectations and bronze validation
    #1.Table level expectations.
    gx_bronze.expect_table_row_count_to_be_between(min_value=1,max_value=1000)
    gx_bronze.expect_table_column_count_to_equal(14)
    gx_bronze.expect_table_columns_to_match_ordered_list(BRONZE_EXPECTED_COLUMNS)   
    
    #2.Column-level Expectations
    gx_bronze.expect_column_values_to_be_in_set("city",value_set=["London", "New York", "Tokyo", "Sydney", "Mumbai"])
    # Not-null checks
    for col in BRONZE_NOT_NULL_COLUMNS:
        gx_bronze.expect_column_values_to_not_be_null(col)
    # Exist-column checks
    for col in BRONZE_EXIST_COLUMNS:
        gx_bronze.expect_column_to_exist(col)
    # Range checks
    for col, rng in BRONZE_RANGE_CHECKS.items():
        gx_bronze.expect_column_values_to_be_between(col, min_value=rng.get("min_value"), max_value=rng.get("max_value"))
    # Regex checks
    for col, rx in BRONZE_REGEX_CHECKS.items():
        gx_bronze.expect_column_values_to_match_regex(col, rx)
    print("BRONZE_REGEX_CHECKS done")
    # Datatype checks
    for col, type in BRONZE_DATATYPE_CHECKS.items():
        gx_bronze.expect_column_values_to_be_of_type(col, type)
    print("BRONZE_DATATYPE_CHECKS done")
    #4.Datetime Format checks
    for col, format in BRONZE_DATETIME_CHECKS.items():
        gx_bronze.expect_column_values_to_match_strftime_format(col,format)
    print("BRONZE_DATETIME_CHECKS done")
    # Validate the bronze dataset.
    bronze_results = gx_bronze.validate()

    # Show results
    show_failed_expectations(bronze_results)
    bronzeSummary = extract_stats("Bronze", bronze_results)
    # Convert GX results to dict first
    bronze_json = bronze_results.to_json_dict()
    write_test_summary_on_doc("Bronze",bronzeSummary, bronze_json)

except Exception as e:
    print(f"Error occurred while reading data from bronze path: {str(e)}")
    logger.error(f"Error occurred while reading data from bronze path: {str(e)}")


# ##### **Perform Hourly Aggregation In Silver Layer**

# In[8]:


from pyspark.sql.functions import col
silver_df = (
         bronze_df
        .withColumn("event_time", to_timestamp("timestamp"))
        .groupBy(
            "city",
            window(col("event_time"), "1 hour")
        )
        .agg(
            avg("temperature").alias("avg_temperature"),
            max("temperature").alias("max_temperature"),
            avg("humidity").alias("avg_humidity"),
            avg("windspeed").alias("avg_windspeed"),
            count("*").alias("record_count")
        )
        .select(
            col("city"),
            col("window.start").alias("hour_start"),
            col("window.end").alias("hour_end"),
            "avg_temperature", "avg_humidity", "avg_windspeed", "max_temperature","record_count"
        )
    )
silver_df = silver_df.withColumn("avg_temperature", round(silver_df["avg_temperature"], 2))
silver_df = silver_df.withColumn("avg_windspeed", round(silver_df["avg_windspeed"], 2))
silver_df = silver_df.withColumn("avg_humidity", round(silver_df["avg_humidity"], 2))


# ##### **Silver Layer Validation With GX**

# In[9]:


# Expectation parameters (example)
SILVER_NOT_NULL_COLUMNS =  ["city","hour_start", "hour_end", "avg_temperature","avg_humidity","avg_windspeed",
                    "max_temperature","record_count"]
SILVER_EXIST_COLUMNS =  ["city","hour_start", "hour_end", "avg_temperature","avg_humidity","avg_windspeed",
                    "max_temperature","record_count"]
SILVER_EXPECTED_COLUMNS =  ["city","hour_start", "hour_end", "avg_temperature","avg_humidity","avg_windspeed",
                    "max_temperature","record_count"]
SILVER_UNIQUE_COMPOUND_COLUMNS = [["hour_start","city"],["hour_end","city"]]
SILVER_DATATYPE_CHECKS = {"city":"StringType", "hour_start":"TimestampType", "hour_end":"TimestampType", "avg_temperature":"DoubleType",
"avg_humidity":"DoubleType", "avg_windspeed":"DoubleType", "max_temperature":"DoubleType", "record_count":"LongType"}
SILVER_RANGE_CHECKS = {"max_temperature": {"min_value": 0, "max_value":50}}

# 4️⃣ Convert to GX Dataset
gx_silver = SparkDFDataset(silver_df)
# Define expectations and silver validation.
gx_silver.expect_table_columns_to_match_ordered_list(SILVER_EXPECTED_COLUMNS) 
# Not-null checks
for col in SILVER_NOT_NULL_COLUMNS:
    gx_silver.expect_column_values_to_not_be_null(col)
for col in SILVER_EXIST_COLUMNS:
    gx_silver.expect_column_to_exist(col)
# Unique checks
for col_list in SILVER_UNIQUE_COMPOUND_COLUMNS:
    gx_silver.expect_compound_columns_to_be_unique(col_list)
# Range checks
for col, rng in SILVER_RANGE_CHECKS.items():
    gx_silver.expect_column_values_to_be_between(col, min_value=rng.get("min_value"), max_value=rng.get("max_value"))
# Datatype checks
for col, type in SILVER_DATATYPE_CHECKS.items():
    gx_silver.expect_column_values_to_be_of_type(col, type)

# Validate the silver dataset.
silver_results = gx_silver.validate()
show_failed_expectations(silver_results)
silverSummary = extract_stats("Silver", silver_results)
# Convert GX results to dict first
silver_json = silver_results.to_json_dict()
write_test_summary_on_doc("Silver",silverSummary, silver_json)


# ##### **Merge Incremental Aggregation into Silver Table**

# In[10]:


silver_path = "Tables/silver/weather_hourly"
if DeltaTable.isDeltaTable(spark, silver_path):
    silver_tbl = DeltaTable.forPath(spark, silver_path)
    silver_tbl.alias("tgt").merge(
        silver_df.alias("src"),
        "tgt.city = src.city AND tgt.hour_start = src.hour_start"
    ).whenMatchedUpdate(set={
        "avg_temperature": "src.avg_temperature",
        "max_temperature": "src.max_temperature",
        "avg_humidity": "src.avg_humidity",
        "avg_windspeed": "src.avg_windspeed",
        "record_count": "src.record_count"
    }).whenNotMatchedInsertAll().execute()
else:
    try:
        logger.info("Performing initial data load on silver path.")
        silver_df.write.format("delta").mode("overwrite").save(silver_path)
    except Exception as e:
        logger.error(f"Error occurred while writing data from silver path: {str(e)}")


# In[11]:


# Optimize table for faster read and merge performance
#spark.sql(f"OPTIMIZE delta.`{silver_path}`ZORDER BY (city)")
try:
    spark.read.format("delta").load(silver_path).orderBy(asc("city"), desc("hour_start")).show(truncate=False)
except Exception as e:
    logger.error(f"Error occurred while reading data from silver path: {str(e)}")


# ##### **Read and Inspect Logs**

# In[12]:


# Read today’s logs
#logger.read_logs()
# Read ERROR logs only
#logger.read_logs(level_filter="ERROR")
# Show last 5 log lines
#logger.read_logs(tail=5)
# Read logs from a specific date (e.g., yesterday)
#logger.read_logs(date="2025-11-06")
# List all log files available
#logger.list_log_files()

